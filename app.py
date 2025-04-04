from flask import Flask, request, jsonify, send_file, render_template, render_template_string
from werkzeug.utils import secure_filename
import os
import util

from docx import Document

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

chat_history = {}

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/check_db/<db_name>', methods=['GET'])
def check_db(db_name):
    if db_name in chat_history.keys():
        return jsonify({'exists': True}), 200
    else:
        return jsonify({'error': 'DB not found'}), 404


@app.route('/ingest', methods=['POST'])
def ingest():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']


    db_name = secure_filename(util.get_unique_filename())
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], db_name)

    if file.filename.endswith('.txt'):
        content = file.read().decode('utf-8')
        doc = SimpleDocTemplate(file_path, pagesize=LETTER)
        styles = getSampleStyleSheet()
        content_paragraphs = [Paragraph(line, styles['Normal']) for line in content.splitlines()]

        doc.build(content_paragraphs)

    elif file.filename.endswith('.docx'):
        doc = Document(file)
        pdf_content = []
        
        for para in doc.paragraphs:
            pdf_content.append(para.text)

        pdf_doc = SimpleDocTemplate(file_path, pagesize=LETTER)
        styles = getSampleStyleSheet()
        content_paragraphs = [Paragraph(line, styles['Normal']) for line in pdf_content]
        
        pdf_doc.build(content_paragraphs)

    elif file.filename.endswith('.pdf') :
        file.save(file_path)

    else : return jsonify({'error': 'Unsupported File Format', 'db_name': '0'})

    util.ingest(file_path, db_name)
    
    return jsonify({'message': 'File ingested successfully', 'db_name': db_name})


@app.route('/query', methods=['POST'])
def query():
    data = request.json
    query_text = data['query']
    db_name = data['db_name']

    if db_name not in chat_history:
        chat_history[db_name] = []

    recent_history = chat_history[db_name][-5:]

    history_context = "\n".join([f"User: {q}\nBot: {r}" for q, r in recent_history])

    cont_aware_query = util.context_aware_query(history_context, query_text) if len(recent_history) > 0 else query_text
    print("Context Aware query:", cont_aware_query)

    try:
        response, sources_with_pages = util.query_rag(cont_aware_query, db_name)
        if sources_with_pages == []:
            response, sources_with_pages = util.query_rag(query_text, db_name)
        

        chat_history[db_name].append((query_text, response))

        return jsonify({'response': response, 'sources': sources_with_pages})

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/get_history/<db_name>', methods=['GET'])
def get_history(db_name):
    if db_name in chat_history:
        history = chat_history[db_name]
        return jsonify({'history': history})
    else:
        return jsonify({'error': 'No history found for this db_name'}), 404



@app.route('/preview/<db_name>')
def preview(db_name):
    if db_name == '0':
        return render_template('default-preview.html')
    file_path = os.path.join(UPLOAD_FOLDER, db_name)
    ext = db_name.split('.')[-1]
    print(file_path)
    if ext == 'pdf':
        return send_file(file_path, mimetype='application/pdf')

    elif ext == 'docx' or ext == 'doc':
        doc = Document(file_path)
        html_content = "<html><body>"
        
        for para in doc.paragraphs:
            html_content += f"<p>{para.text}</p>"
        
        html_content += "</body></html>"

        return render_template_string(html_content)
    
    elif ext == 'txt':
        return send_file(file_path, mimetype='text/plain')

    else:
        return jsonify({'error': 'Unsupported file type'}), 400

from reportlab.lib.pagesizes import LETTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

@app.route('/export_chat/<db_name>', methods=['GET'])
def export_chat(db_name):

    if db_name not in chat_history:
        return jsonify({'error': 'No history found for this db_name'}), 404

    history = chat_history[db_name]

    pdf_filename = f"chat_history_{db_name}.pdf"
    pdf_filepath = os.path.join(app.config['UPLOAD_FOLDER'], pdf_filename)

    doc = SimpleDocTemplate(pdf_filepath, pagesize=LETTER)
    styles = getSampleStyleSheet()
    content = []

    content.append(Paragraph(f"Chat History: {db_name}", styles['Title']))
    content.append(Spacer(1, 12))

    for i, (user, bot) in enumerate(history):
        content.append(Paragraph(f"Q{i+1}: {user}", styles['Normal']))
        content.append(Spacer(1, 6))
        content.append(Paragraph(f"A{i+1}: {bot}", styles['Normal']))
        content.append(Spacer(1, 12))

    doc.build(content)

    return send_file(pdf_filepath, as_attachment=True, download_name=pdf_filename)


if __name__ == '__main__':
    app.run(debug=True)
